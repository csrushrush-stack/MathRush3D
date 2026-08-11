from __future__ import annotations

import os
import shutil
import subprocess
from zipfile import ZipFile
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\Acer\Documents\MathsRush3D")
WORK = ROOT / "assignment_work" / "math_rush_report"
FIGURES = WORK / "figures"
OUTPUT = ROOT / "deliverables" / "Math_Rush_3D_Application_Development_Report_PRE-HUMANIZER.docx"
TEMPLATE = Path(r"C:\Users\Acer\.codex\skills\artifact-template-vincent-assignment-document\assets\reference.docx")
TEMPLATE_LOGO = WORK / "cosmo_logo.png"

NAVY = "0B2545"
BLUE = "2E74B5"
CYAN = "0E7490"
GOLD = "D59B14"
INK = "1D2733"
MUTED = "5C6B7A"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
PALE_CYAN = "E6F4F6"
PALE_GOLD = "FFF7DF"
RED = "9B1C1C"
GREEN = "18794E"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int],
              font_size: float = 12) -> object:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = rgb(WHITE)
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    set_repeat_table_header(table.rows[0])
    for row_idx, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        if row_idx % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(text))
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.color.rgb = rgb(INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def set_run_font(run, size=12, color=INK, bold=False, italic=False, name="Times New Roman",
                 underline=False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             after=6, keep=False) -> object:
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    if keep:
        p.paragraph_format.keep_with_next = True
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str, level=0) -> object:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc: Document, text: str) -> object:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc: Document, label: str, text: str, fill=PALE_GOLD, accent=GOLD) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=10, color=MUTED, italic=True)


def add_picture(doc: Document, path: Path, caption: str, alt_text: str, width=6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    inline = shape._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", alt_text)
    add_caption(doc, caption)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, code: str, display: str = "") -> None:
    begin_run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_char)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(display)
    set_run_font(result_run, size=9, color=MUTED)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Table of contents will update in Microsoft Word.")


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def h1(doc: Document, text: str, new_page=True) -> None:
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.page_break_before = new_page


def h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def dot_render(name: str, source: str) -> Path:
    dot_file = FIGURES / f"{name}.dot"
    png_file = FIGURES / f"{name}.png"
    dot_file.write_text(source, encoding="utf-8")
    dot_exe = shutil.which("dot") or r"C:\Users\Acer\scoop\apps\graphviz\current\bin\dot.exe"
    subprocess.run([dot_exe, "-Tpng", "-Gdpi=180", str(dot_file), "-o", str(png_file)], check=True)
    return png_file


def make_diagrams() -> dict[str, Path]:
    base = 'graph [bgcolor="white", pad="0.35", nodesep="0.62", ranksep="0.78", splines=ortho, outputorder="edgesfirst", overlap=false]; node [fontname="Arial", fontsize=12, color="#0B2545", fontcolor="#1D2733", penwidth=1.5, margin="0.16,0.10"]; edge [fontname="Arial", fontsize=10, color="#5C6B7A", fontcolor="#5C6B7A", penwidth=1.3, arrowsize=0.75];'
    diagrams = {}
    diagrams["use_case"] = dot_render("use_case", f'''digraph G {{ {base}
      rankdir=TB; graph [splines=polyline, nodesep="0.48", ranksep="0.72"]; node [shape=ellipse, style="filled", fillcolor="#E8EEF5"];
      title [shape=plaintext, label="Math Rush 3D use cases", fontname="Arial", fontsize=16, fontcolor="#0B2545"];
      guest [shape=box, label="Guest", fillcolor="#FFF7DF"];
      player [shape=box, label="Authenticated Player", fillcolor="#E6F4F6"];
      admin [shape=box, label="Administrator\n(planned)", style="filled,dashed", fillcolor="#F4F6F9"];
      register [label="Register or log in"];
      choose [label="Choose difficulty\nand level"];
      play [label="Play a run"];
      board [label="View leaderboard\nand statistics"];
      skins [label="Buy or equip skins"];
      manage [label="Manage users, content\nand reports (planned)", style="filled,dashed"];
      gates [label="Solve gate choices"];
      obstacles [label="Avoid obstacles"];
      boss [label="Defeat boss"];
      results [label="Save score and\nprogress"];
      title -> player [style=invis, weight=30];
      {{rank=same; guest; player; admin;}}
      guest -> player -> admin [style=invis, weight=30];
      {{rank=same; register; choose; manage;}}
      register -> choose -> manage [style=invis, weight=30];
      {{rank=same; play; board; skins;}}
      play -> board -> skins [style=invis, weight=30];
      {{rank=same; gates; obstacles; boss; results;}}
      gates -> obstacles -> boss -> results [style=invis, weight=30];
      guest -> register; guest -> play;
      player -> choose; player -> play; player -> board; player -> skins;
      play -> gates [style=dashed]; play -> obstacles [style=dashed]; play -> boss [style=dashed]; play -> results [style=dashed];
      admin -> manage [style=dashed];
    }}''')

    diagrams["erd"] = dot_render("erd", f'''digraph G {{ {base}
      rankdir=LR; graph [nodesep="0.34", ranksep="0.95"]; node [shape=record, style="filled", fillcolor="#F4F6F9", fontsize=10.5, margin="0.12,0.08"];
      players [label="players|PK id\ldevice_id\ldisplay_name\lselected_skin FK\l"];
      accounts [label="player_accounts (0..1 per player)|PK/FK player_id\lemail\lpassword_hash\l"];
      sessions [label="auth_sessions (many per player)|PK id\lplayer_id FK\ltoken_hash\lexpires_at\l"];
      progress [label="player_progress (1 per player)|PK/FK player_id\lcoins\lbest_score\llevels completed\l"];
      settings [label="player_settings (1 per player)|PK/FK player_id\lmusic / SFX\lvibration\lreduced_effects\l"];
      skins [label="skins|PK id\lname\lprice\lis_available\l"];
      player_skins [label="player_skins (many-to-many)|PK/FK player_id\lPK/FK skin_id\lacquired_at\l"];
      runs [label="game_runs (many per player)|PK id\lplayer_id FK\ldifficulty / level\lscore / status\lmultiplier / stars\l"];
      gates [label="gate_choices (many per run)|PK id\lrun_id FK\lexpressions\lchosen / optimal delta\l"];
      obstacles [label="obstacle_events (many per run)|PK id\lrun_id FK\ltype / outcome\ldamage\l"];
      achievements [label="achievements|PK id\lname\ltarget / reward\l"];
      player_ach [label="player_achievements (many-to-many)|PK/FK player_id\lPK/FK achievement_id\lprogress\l"];
      {{rank=same; skins; players; achievements;}}
      {{rank=same; accounts; sessions; progress; settings; runs; player_skins; player_ach;}}
      {{rank=same; gates; obstacles;}}
      players -> accounts; players -> sessions; players -> progress; players -> settings; players -> runs;
      skins -> players [style=dashed]; players -> player_skins; skins -> player_skins;
      runs -> gates; runs -> obstacles; players -> player_ach; achievements -> player_ach;
    }}''')

    diagrams["architecture"] = dot_render("architecture", f'''digraph G {{ {base}
      rankdir=TB; graph [ranksep="0.72"]; node [shape=box, style="rounded,filled", fillcolor="#E8EEF5", width=2.25, height=0.72];
      subgraph cluster_client {{ label="Presentation tier"; color="#2E74B5"; style="rounded"; margin=22;
        browser [label="Web browser\nReact + Tailwind"];
        android [label="Android APK\nCapacitor WebView"];
        core [label="Shared client core\nR3F + Three.js + Zustand", fillcolor="#E6F4F6", width=3.0];
        {{rank=same; browser; android;}}
        browser -> core; android -> core;
      }}
      api_link [label="HTTPS / JSON\nbearer session token", shape=box, style="rounded,filled", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", width=2.35, height=0.54];
      subgraph cluster_api {{ label="Application tier"; color="#0E7490"; style="rounded"; margin=22;
        api [label="Express REST API\nTypeScript + Zod", fillcolor="#E6F4F6", width=2.8];
        auth [label="Authentication\nSessions + bcrypt"];
        rules [label="Reward rules\nand validation"];
        {{rank=same; auth; rules;}}
        api -> auth; api -> rules;
      }}
      sql_link [label="Parameterised SQL", shape=box, style="rounded,filled", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", width=2.0, height=0.45];
      subgraph cluster_data {{ label="Data tier"; color="#D59B14"; style="rounded"; margin=22;
        pg [label="PostgreSQL\nNeon production", width=2.8];
      }}
      core -> api_link -> api;
      api -> sql_link -> pg;
    }}''')

    diagrams["dfd0"] = dot_render("dfd0", f'''digraph G {{ {base}
      rankdir=TB; graph [ranksep="0.58"]; node [shape=box, style="rounded,filled", fillcolor="#FFF7DF"];
      player [label="Player", width=1.5];
      player_flow [label="To system:\ncredentials, controls, settings and gate choices\n\nTo player:\n3D gameplay, results, progress and leaderboard", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", width=3.3];
      system [shape=circle, label="0\nMath Rush 3D", fillcolor="#E6F4F6", width=2.0];
      database_flow [label="To database:\naccounts, runs, events and progress updates\n\nTo system:\nprofiles, unlocks, rankings and statistics", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", width=3.2];
      db [shape=cylinder, label="PostgreSQL\nGame database", fillcolor="#E8EEF5", width=2.1];
      admin [label="Administrator\n(planned)", style="rounded,filled,dashed"];
      admin_flow [label="Management requests\nand audit reports", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", style="rounded,filled,dashed"];
      player -> player_flow [dir=both]; player_flow -> system [dir=both];
      system -> database_flow [dir=both]; database_flow -> db [dir=both];
      {{rank=same; system; admin_flow; admin;}}
      system -> admin_flow [dir=both, style=dashed, constraint=false]; admin_flow -> admin [dir=both, style=dashed, constraint=false];
    }}''')

    diagrams["dfd1"] = dot_render("dfd1", f'''digraph G {{ {base}
      rankdir=TB; graph [ranksep="0.48"]; node [shape=box, style="rounded,filled", fillcolor="#E6F4F6", width=2.65, height=0.60];
      player_in [label="Player\n(input)", fillcolor="#FFF7DF", width=1.7];
      f_credentials [label="Email, password and device identity", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      auth [label="1.0  Authenticate player"];
      f_session [label="Verified session and player profile", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      game [label="2.0  Run game locally"];
      f_run [label="Run summary and event evidence", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      submit [label="3.0  Validate and save run"];
      f_rewards [label="Trusted score and reward values", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      prog [label="4.0  Update progress and skins"];
      f_progress [label="Progress, settings and unlock state", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      report [label="5.0  Produce leaderboard and statistics"];
      f_rankings [label="Rankings and player statistics", fillcolor="#FFFFFF", color="#B8C6D5", fontcolor="#5C6B7A", height=0.42];
      player_out [label="Player\n(results)", fillcolor="#FFF7DF", width=1.7];
      d1 [shape=cylinder, label="D1 Accounts\nand sessions", fillcolor="#E8EEF5"];
      d2 [shape=cylinder, label="D2 Runs and\nevent history", fillcolor="#E8EEF5"];
      d3 [shape=cylinder, label="D3 Progress, settings\nand unlocks", fillcolor="#E8EEF5"];
      player_in -> f_credentials -> auth -> f_session -> game -> f_run -> submit -> f_rewards -> prog -> f_progress -> report -> f_rankings -> player_out;
      {{rank=same; auth; d1;}} auth -> d1 [dir=both, constraint=false];
      {{rank=same; submit; d2;}} submit -> d2 [dir=both, constraint=false];
      {{rank=same; prog; d3;}} prog -> d3 [dir=both, constraint=false];
      d2 -> report [style=invis, constraint=false]; d3 -> report [style=invis, constraint=false];
    }}''')

    diagrams["flow"] = dot_render("gameplay_flow", f'''digraph G {{ {base}
      rankdir=TB; graph [ranksep="0.50"]; node [shape=box, style="rounded,filled", fillcolor="#E8EEF5", width=2.75, height=0.60];
      start [shape=oval, label="Launch application", fillcolor="#FFF7DF"];
      auth [label="Register or log in"]; home [label="Home screen"]; select [label="Select unlocked difficulty and level"]; run [label="Auto-run and steer crowd"];
      gate [shape=diamond, label="Math gate reached?", fillcolor="#E6F4F6", width=2.6];
      apply_gate [label="Apply selected answer"];
      obs [shape=diamond, label="Obstacle or enemy contact?", fillcolor="#E6F4F6", width=3.1];
      apply_hit [label="Apply collision to affected members"];
      alive [shape=diamond, label="Crowd remaining?", fillcolor="#E6F4F6", width=2.65];
      boss_reached [shape=diamond, label="Boss area reached?", fillcolor="#E6F4F6", width=2.8];
      continue_run [label="Continue the course"];
      meter [label="Stop the boss boost meter"];
      boss [shape=diamond, label="Boss defeated?", fillcolor="#E6F4F6", width=2.55];
      bonus [label="Run through bonus multipliers\nuntil the crowd is spent"];
      win [label="Win: save score, stars, coins,\nprogress and next level"];
      lose [label="Game over: save run result", fillcolor="#FDECEC"];
      end [shape=oval, label="Replay, next level or home", fillcolor="#FFF7DF"];
      yes_gate [label="YES", shape=box, width=0.72, height=0.34, fillcolor="#E9F7F0", color="#8BC5A8", fontcolor="#18794E"];
      no_gate [label="NO", shape=box, width=0.72, height=0.34, fillcolor="#F4F6F9", color="#B8C6D5", fontcolor="#5C6B7A"];
      yes_obs [label="YES", shape=box, width=0.72, height=0.34, fillcolor="#FFF7DF", color="#D9BD70", fontcolor="#8A6508"];
      no_obs [label="NO", shape=box, width=0.72, height=0.34, fillcolor="#F4F6F9", color="#B8C6D5", fontcolor="#5C6B7A"];
      yes_alive [label="YES", shape=box, width=0.72, height=0.34, fillcolor="#E9F7F0", color="#8BC5A8", fontcolor="#18794E"];
      no_alive [label="NO", shape=box, width=0.72, height=0.34, fillcolor="#FDECEC", color="#D7A0A0", fontcolor="#9B1C1C"];
      yes_boss_area [label="YES", shape=box, width=0.72, height=0.34, fillcolor="#E9F7F0", color="#8BC5A8", fontcolor="#18794E"];
      no_boss_area [label="NO", shape=box, width=0.72, height=0.34, fillcolor="#F4F6F9", color="#B8C6D5", fontcolor="#5C6B7A"];
      yes_boss [label="YES", shape=box, width=0.72, height=0.34, fillcolor="#E9F7F0", color="#8BC5A8", fontcolor="#18794E"];
      no_boss [label="NO", shape=box, width=0.72, height=0.34, fillcolor="#FDECEC", color="#D7A0A0", fontcolor="#9B1C1C"];
      start -> auth -> home -> select -> run -> gate;
      gate -> yes_gate -> apply_gate -> obs;
      gate -> no_gate -> obs;
      obs -> yes_obs -> apply_hit -> alive;
      obs -> no_obs -> alive;
      alive -> no_alive -> lose;
      alive -> yes_alive -> boss_reached;
      boss_reached -> no_boss_area -> continue_run;
      continue_run -> run [constraint=false];
      boss_reached -> yes_boss_area -> meter -> boss;
      boss -> no_boss -> lose;
      boss -> yes_boss -> bonus -> win -> end;
      lose -> end;
    }}''')
    diagrams["wireframes"] = make_wireframes()
    return diagrams


def font(size: int, bold=False):
    candidates = [
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def make_wireframes() -> Path:
    w, h = 1600, 1680
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font, h_font, body_font, small_font = font(34, True), font(22, True), font(17), font(14)
    d.text((55, 30), "Math Rush 3D: low-fidelity mobile wireframes", fill="#0B2545", font=title_font)

    panels = [
        (110, 105, "1. Authentication"),
        (880, 105, "2. Home / Level Select"),
        (110, 855, "3. Gameplay"),
        (880, 855, "4. Results"),
    ]
    for x, y, heading in panels:
        d.rounded_rectangle((x, y, x + 610, y + 660), radius=35, fill="#F4F6F9", outline="#0B2545", width=4)
        d.text((x + 22, y + 22), heading, fill="#0B2545", font=h_font)
    # Authentication
    x, y = 110, 105
    d.rounded_rectangle((x + 90, y + 85, x + 520, y + 165), radius=18, fill="#E8EEF5", outline="#2E74B5", width=3)
    d.text((x + 228, y + 112), "MATH RUSH", fill="#2E74B5", font=h_font)
    for idx, label in enumerate(["Runner name", "Email", "Password"]):
        yy = y + 205 + idx * 100
        d.text((x + 75, yy), label, fill="#5C6B7A", font=small_font)
        d.rounded_rectangle((x + 75, yy + 24, x + 535, yy + 75), radius=12, fill="white", outline="#9AA8B5", width=2)
    d.rounded_rectangle((x + 75, y + 535, x + 535, y + 600), radius=16, fill="#0E7490")
    d.text((x + 274, y + 555), "ENTER", fill="white", font=h_font)
    d.text((x + 205, y + 620), "Scrollable on small screens", fill="#5C6B7A", font=small_font)
    # Home
    x, y = 880, 105
    d.text((x + 210, y + 86), "Welcome, Runner", fill="#0B2545", font=h_font)
    d.rounded_rectangle((x + 75, y + 130, x + 535, y + 200), radius=18, fill="#E6F4F6", outline="#0E7490", width=2)
    d.text((x + 225, y + 153), "Coins / stars / rank", fill="#0E7490", font=body_font)
    d.text((x + 75, y + 225), "Difficulty", fill="#5C6B7A", font=small_font)
    for idx, label in enumerate(["EASY", "MEDIUM", "HARD", "EXPERT"]):
        yy = y + 255 + idx * 62
        fill = "#2E74B5" if idx == 0 else "#DDE3E9"
        color = "white" if idx == 0 else "#5C6B7A"
        d.rounded_rectangle((x + 75, yy, x + 535, yy + 48), radius=12, fill=fill)
        d.text((x + 275, yy + 13), label, fill=color, font=body_font)
    d.text((x + 75, y + 520), "Levels  1  2  3  4  5", fill="#0B2545", font=body_font)
    d.rounded_rectangle((x + 75, y + 560, x + 535, y + 615), radius=16, fill="#D59B14")
    d.text((x + 280, y + 576), "PLAY", fill="white", font=h_font)
    # Gameplay
    x, y = 110, 855
    d.text((x + 75, y + 85), "Crowd 42                                     Pause", fill="#0B2545", font=body_font)
    d.rectangle((x + 155, y + 120, x + 455, y + 590), fill="#DCE6EC", outline="#5C6B7A", width=2)
    d.rounded_rectangle((x + 180, y + 190, x + 280, y + 265), radius=8, fill="#0E7490")
    d.rounded_rectangle((x + 330, y + 190, x + 430, y + 265), radius=8, fill="#D59B14")
    d.text((x + 212, y + 214), "8+7", fill="white", font=small_font)
    d.text((x + 362, y + 214), "3x4", fill="white", font=small_font)
    d.ellipse((x + 286, y + 425, x + 326, y + 465), fill="#2E74B5")
    for dx, dy in [(-45, 34), (0, 42), (45, 34), (-72, 78), (-25, 84), (25, 84), (72, 78)]:
        d.ellipse((x + 286 + dx, y + 425 + dy, x + 316 + dx, y + 455 + dy), fill="#2E74B5")
    d.text((x + 245, y + 610), "Drag left / right", fill="#5C6B7A", font=body_font)
    # Results
    x, y = 880, 855
    d.text((x + 255, y + 86), "VICTORY", fill="#18794E", font=h_font)
    d.text((x + 252, y + 125), "3 STARS", fill="#D59B14", font=h_font)
    d.rounded_rectangle((x + 75, y + 175, x + 535, y + 365), radius=18, fill="#E8EEF5", outline="#2E74B5", width=2)
    for idx, label in enumerate(["Score       12,450", "Multiplier      x7", "Coins          +101"]):
        d.text((x + 145, y + 215 + idx * 48), label, fill="#0B2545", font=body_font)
    for idx, (label, fill) in enumerate([("NEXT LEVEL", "#0E7490"), ("REPLAY", "#2E74B5"), ("HOME", "#5C6B7A")]):
        yy = y + 400 + idx * 68
        d.rounded_rectangle((x + 75, yy, x + 535, yy + 52), radius=14, fill=fill)
        d.text((x + 255, yy + 15), label, fill="white", font=body_font)
    d.text((170, 1600), "Four principal mobile states are shown at a readable size. Final annotated Figma evidence remains to be supplied.", fill="#5C6B7A", font=body_font)
    path = FIGURES / "wireframes.png"
    img.save(path, dpi=(180, 180))
    return path


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for style_name, size, color, before, after in (
        ("Heading 1", 14, INK, 18, 10),
        ("Heading 2", 12, INK, 12, 6),
        ("Heading 3", 12, INK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.underline = style_name in ("Heading 1", "Heading 2")
        style.font.italic = style_name == "Heading 3"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name, left, first in (
        ("List Bullet", 0.375, -0.194),
        ("List Bullet 2", 0.625, -0.194),
        ("List Number", 0.375, -0.194),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)

    if "Reference Entry" not in styles:
        ref = styles.add_style("Reference Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference Entry"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    ref.font.size = Pt(12)
    ref.paragraph_format.left_indent = Inches(0.3)
    ref.paragraph_format.first_line_indent = Inches(-0.3)
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.line_spacing = 1.5


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.70)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.35)
    header_table.columns[1].width = Inches(2.55)
    for cell in header_table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "nil")
            borders.append(edge)
        tc_pr.append(borders)
    left = header_table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    r = left.add_run("UNIT 22: APPLICATION DEVELOPMENT")
    set_run_font(r, size=9, color=MUTED, bold=True)
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    r = right.add_run("INDIVIDUAL ASSIGNMENT")
    set_run_font(r, size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("HNDC18\t")
    set_run_font(r, size=9, color=MUTED)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25))
    add_field(p, "PAGE", "1")


def extract_template_logo() -> Path:
    if not TEMPLATE_LOGO.exists():
        with ZipFile(TEMPLATE) as archive:
            TEMPLATE_LOGO.write_bytes(archive.read("word/media/image1.png"))
    return TEMPLATE_LOGO



def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    p.add_run().add_picture(str(extract_template_logo()), width=Inches(2.25))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIT 22: APPLICATION DEVELOPMENT")
    set_run_font(r, size=14, color=INK)
    p.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(MATH RUSH 3D REPORT)")
    set_run_font(r, size=14, color=INK)
    p.paragraph_format.space_after = Pt(48)

    for value in ("VINCENT ESCANDALLO CASTILLO", "TJ55730"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(value)
        set_run_font(r, size=14, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REPORT SUBMITTED FOR THE")
    set_run_font(r, size=14, color=INK)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PEARSON BTEC HIGHER NATIONAL DIPLOMA IN COMPUTING")
    set_run_font(r, size=14, color=INK)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COSMOPOLITAN COLLEGE OF COMMERCE AND TECHNOLOGY")
    set_run_font(r, size=14, color=INK)
    page_break(doc)


def add_front_matter(doc: Document) -> None:
    h1(doc, "Report scope and evidence status", new_page=False)
    add_callout(doc, "Evidence status", "This report addresses the supplied project guideline and assignment brief using the inspected Math Rush 3D implementation. Missing peer-review, device-performance and administration evidence is stated explicitly and is not presented as completed work.", fill=PALE_GOLD, accent=GOLD)
    h2(doc, "Revision record")
    add_table(doc, ["Version", "Date", "Purpose", "Status"], [
        ("0.1", "21 July 2026", "Initial report based on currently supplied guidelines", "Draft"),
        ("0.2", "29 July 2026", "Reformatted to the requested Times New Roman academic style and expanded traceability, performance and ethics sections", "Revised draft"),
        ("1.0", "4 August 2026", "Final report review, current build verification and diagram layout correction", "Final report"),
    ], [1050, 1650, 4700, 1960])
    h1(doc, "Table of contents")
    add_toc(doc)


def add_executive_summary(doc: Document) -> None:
    h1(doc, "Executive summary", new_page=False)
    add_body(doc, "Math Rush 3D is a mobile-first three-dimensional crowd-running game that combines fast arcade play with short arithmetic decisions. The project responds to a practical engagement problem: learners often experience basic arithmetic practice as repetitive and disconnected from immediate consequences. In Math Rush 3D, addition, subtraction, multiplication and division are placed directly in the player's route. Correct or advantageous decisions visibly change the size of the running crowd, which then determines whether the player can survive obstacles, defeat an end-level boss and reach a higher bonus multiplier.")
    add_body(doc, "The implemented product is delivered as both a responsive web application and an Android package. The presentation tier uses React, TypeScript, Tailwind CSS, Three.js and React Three Fiber. Zustand coordinates interface and run state without forcing frame-by-frame React updates. The application tier is currently an Express REST API written in TypeScript, with Zod validation, rate limiting, security headers and session-based authentication. PostgreSQL stores accounts, settings, progression, run summaries, gate choices, obstacle events, skins, achievements and leaderboard data. Production services use Render and Neon, while local development can use Docker.")
    add_body(doc, "A typical session begins with registration or login, followed by difficulty and level selection. The player steers a clustered crowd through ten pairs of arithmetic gates and difficulty-specific obstacles. Enemy crowds and hazards remove individual runners. Immediately before the boss, the player stops a timing meter to earn a crowd boost. The boss battle consumes crowd members according to the level balance. If successful, the survivors enter a multiplier lane, where each bonus stage exchanges crowd members for bonus points until the crowd is exhausted. The result screen awards a score, coins, stars, level progression and a next-level option.")
    add_body(doc, "The current implementation demonstrates a substantial functional baseline: four difficulty bands with five levels each, responsive controls, GPU-instanced crowds, collision tests based on movement-plane crossing, a balanced boss calculation, skins, persistent accounts and an online leaderboard. On 4 August 2026, all 30 automated tests passed, the browser and server TypeScript checks completed successfully, and Vite produced a production build. These checks support the internal consistency of the balancing, formation, level-progression, speed and boss-meter utilities, but they do not replace device testing, accessibility review, penetration testing or user acceptance testing.")
    add_body(doc, "Several assignment-compliance items remain open. The brief lists PHP, Python or Java as expected backend technologies and permits Node.js only as middleware or API glue; therefore, the current Node.js/Express backend requires explicit tutor approval or replacement by an accepted service. Role-based access control, an administrator interface, full core-entity CRUD, a formal audit trail, a feedback store and completed peer-review evidence are also not yet implemented. These gaps are documented rather than hidden so that the next development iteration can be planned against the marking criteria.")
    add_body(doc, "The recommended next stage is to obtain written confirmation about backend-language compliance, then implement an administrator role and management workflow, complete remaining CRUD/search/report functions, collect peer feedback using the UAT checklist, and add dated screenshots and GitHub evidence. The design and research in this report provide a traceable starting point for those activities and can be expanded when the remaining guideline is supplied.")


def add_intro(doc: Document) -> None:
    h1(doc, "1. Project context and scope")
    h2(doc, "1.1 Scenario")
    add_body(doc, "The project is framed as a GameForge Studios hackathon entry. The challenge is to create a browser-based game that addresses a real-world problem while demonstrating application-development practice, relational data management and evidence-led evaluation. Math Rush 3D interprets the real-world problem as low engagement with repeated mental-arithmetic practice, particularly when exercises provide little agency or visual feedback.")
    h2(doc, "1.2 Aim")
    add_body(doc, "The aim is to create a responsive, replayable arithmetic game in which a player's mathematical decisions and movement skill jointly determine success, while a secure data service records progress, settings and competitive results across web and Android clients.")
    h2(doc, "1.3 Objectives")
    for item in (
        "Deliver a complete start-to-finish game loop with gates, hazards, enemy crowds, a boss and a finite bonus ending.",
        "Scale arithmetic content, obstacles and movement speed across four difficulty bands and five levels per band.",
        "Maintain fair level balance so that an optimal route remains achievable while poor choices have meaningful consequences.",
        "Provide account registration, login, saved progression, skins, statistics and leaderboards through PostgreSQL.",
        "Support touch, mouse and keyboard interaction on responsive web layouts and Android devices.",
        "Collect evidence through automated checks, manual tests, peer UAT and GitHub history.",
    ):
        add_bullet(doc, item)
    h2(doc, "1.4 Success measures")
    add_table(doc, ["Measure", "Target", "Evidence method"], [
        ("Mathematical correctness", "All generated expressions return the displayed integer result and update the crowd once.", "Automated gate tests plus sampled gameplay evidence."),
        ("Level fairness", "An optimal route can defeat the boss in every published difficulty and level.", "Balance tests and documented perfect-route play tests."),
        ("Mobile usability", "Core controls and authentication remain usable at 360x640 and larger supported viewports.", "Responsive browser matrix and Android UAT."),
        ("Runtime performance", "Gameplay should normally remain at or above 45 FPS on the selected mid-range Android test device.", "Recorded frame-rate and frame-time test runs."),
        ("Persistence reliability", "A completed authenticated run appears once in progress and leaderboard records after reconnection.", "API/database integration test and offline queue test."),
        ("Security baseline", "Invalid input is rejected, passwords are never stored in plain text and unauthorised admin access is denied.", "Validation, database inspection and authorisation tests."),
    ], [2600, 3550, 3210])
    h2(doc, "1.5 Scope boundaries")
    add_body(doc, "The current scope includes a single-player crowd runner, arithmetic gates, local 3D simulation, account-backed persistence and competitive rankings. Real-time multiplayer, payments, advertisements, social messaging, user-generated levels and an iOS package are excluded. The administrator role and formal feedback workflow are required by the brief but are identified as planned work because they are not present in the inspected repository.")


def add_sdd(doc: Document, figs: dict[str, Path]) -> None:
    h1(doc, "2. Software Design Document")
    h2(doc, "2.1 Problem definition")
    add_body(doc, "Conventional arithmetic worksheets can test correctness but rarely sustain voluntary repetition. They also separate the calculation from a meaningful consequence. Math Rush 3D turns arithmetic into an immediate spatial decision: the player reads two expressions, steers towards one answer and sees the crowd change. The design must avoid becoming a guessing game, however. Symbols must remain legible on small screens, level generation must not create impossible encounters, and movement skill must reduce casualties without allowing all hazards to be ignored.")
    h2(doc, "2.2 Stakeholder analysis")
    add_table(doc, ["Stakeholder", "Need or interest", "Influence", "Design response"], [
        ("Learner / player", "Clear rules, fair challenge, responsive steering, visible rewards and saved progress.", "High", "Difficulty progression, readable gates, touch controls, feedback and next-level flow."),
        ("Lecturer / assessor", "Evidence that the solution satisfies the brief and is evaluated honestly.", "High", "Traceable requirements, diagrams, testing records, references and explicit compliance gaps."),
        ("Game developer", "Maintainable modules, reproducible builds and measurable defects.", "High", "TypeScript, separated game systems, shared rules, tests and GitHub workflow."),
        ("Administrator", "Manage users/content and inspect reports.", "Medium", "Planned RBAC, administration dashboard, audit log and CRUD endpoints."),
        ("Hosting/database provider", "Secure, efficient and correctly configured workloads.", "Medium", "Connection pooling, environment variables, rate limits and health checks."),
        ("Parent/educator", "Age-appropriate play and evidence of mathematical value.", "Low/medium", "No chat or payments; arithmetic decisions and progression statistics."),
    ], [1500, 3100, 1050, 3710])

    h2(doc, "2.3 User requirements")
    add_table(doc, ["ID", "Requirement", "Acceptance measure", "Status"], [
        ("UR-01", "A visitor can register and an existing player can log in.", "Valid credentials create a persisted session; invalid input receives a clear error.", "Implemented"),
        ("UR-02", "A player can choose an unlocked difficulty and level.", "Locked stages are visibly unavailable and completing a stage unlocks the next valid stage.", "Implemented"),
        ("UR-03", "A player can steer by touch, mouse or keyboard.", "The crowd moves laterally within track boundaries without page scrolling during play.", "Implemented"),
        ("UR-04", "Gate content contains difficulty-appropriate arithmetic expressions.", "Addition, subtraction, multiplication and division appear at suitable complexity and produce integer outcomes.", "Implemented"),
        ("UR-05", "Individual crowd members can be removed by hazards.", "A hit reduces the crowd by the calculated amount and produces visible feedback.", "Implemented"),
        ("UR-06", "The level has a boss and a finite bonus ending.", "A successful boss fight enters the multiplier lane and finishes at a visible end structure.", "Implemented"),
        ("UR-07", "A player can pause, restart, replay or continue to the next level.", "Controls change the current run without requiring a browser refresh.", "Implemented"),
        ("UR-08", "Progress, scores, settings and skins remain available after sign-in on another client.", "Server-backed profile values reload after authentication.", "Implemented"),
        ("UR-09", "Players can view rankings and personal statistics.", "Leaderboard and statistics panels retrieve current stored data.", "Implemented / verify UAT"),
        ("UR-10", "An administrator can manage users, game content and reports.", "Authorised admin screens support protected CRUD and reporting actions.", "Planned"),
    ], [800, 3250, 3750, 1560])

    h2(doc, "2.4 System requirements")
    h3(doc, "Functional requirements")
    for item in (
        "FR-01: generate exactly ten arithmetic gate pairs for a standard level and retain the chosen and optimal values for run evidence.",
        "FR-02: calculate obstacle, enemy and boss outcomes from a deterministic level definition shared by all scene systems.",
        "FR-03: validate registration, login, player updates and run submissions at the API boundary.",
        "FR-04: use parameterised SQL to create and retrieve account, progression, run, event, skin and leaderboard records.",
        "FR-05: prevent duplicate run submission by treating the client run identifier as unique.",
        "FR-06: support offline/local play and queue a completed run for later synchronisation when the API is unavailable.",
        "FR-07: calculate rewards using a shared rule module so browser and API calculations cannot silently diverge.",
        "FR-08: provide protected administrator CRUD, RBAC, audit reporting, feedback collection and search/filter functions (planned requirement).",
    ):
        add_bullet(doc, item)
    h3(doc, "Non-functional requirements")
    for item in (
        "NFR-01 Performance: target smooth play on mainstream mobile devices by using instanced crowds, capped pixel ratio, reduced-effects settings and ref-based frame data.",
        "NFR-02 Usability: maintain readable arithmetic, large touch targets, safe-area spacing and scrollable authentication forms on small screens.",
        "NFR-03 Reliability: recover gracefully when the API is sleeping or temporarily unreachable; avoid losing a locally completed run.",
        "NFR-04 Security: hash passwords, store session tokens as hashes, use HTTPS in production, apply security headers, rate-limit authentication and keep secrets outside the repository.",
        "NFR-05 Maintainability: separate scene systems, pure game rules, API routes, validation, database access and shared reward logic.",
        "NFR-06 Compatibility: support current Chromium-based browsers and Android 7.0/API 24 or later through Capacitor.",
        "NFR-07 Accessibility: preserve keyboard operation outside gesture-only gameplay, meaningful labels, sufficient contrast and reduced-effects options; formal WCAG testing remains pending.",
    ):
        add_bullet(doc, item)

    h2(doc, "2.5 Assumptions and constraints")
    add_body(doc, "The design assumes an internet connection is available for account synchronisation, but the core run remains local and playable during an API outage. Render's free service may cold-start, so the first request can be slower than subsequent requests. The 3D scene is constrained by mobile GPU and memory limits, making instancing and bounded effects more important than photorealistic assets. The largest academic constraint is that the current Express API may not satisfy the brief's prescribed backend-language list; this must be resolved with the tutor before final submission.")

    h2(doc, "2.6 Risk review and mitigation")
    add_table(doc, ["Risk", "Likelihood", "Impact", "Mitigation / contingency", "Owner"], [
        ("Generated level is impossible even with optimal play", "Medium", "High", "Balance from the same level definition; automated game-balance tests; play-test every difficulty/level seed policy.", "Developer"),
        ("Frame-rate drop with large crowds", "Medium", "High", "GPU instancing, reusable geometry/materials, limited pixel ratio, reduced-effects mode and device profiling.", "Developer"),
        ("API cold start is interpreted as login failure", "High", "Medium", "Visible connecting state, retry/backoff, health check and clearer timeout message; consider paid/always-on hosting later.", "Developer"),
        ("Credential or session compromise", "Low/medium", "High", "bcrypt work factor, token hashing, secure HTTP-only cookie, short-scoped bearer storage, rate limits, HTTPS and session expiry.", "Developer"),
        ("Leaderboard score is manipulated", "Medium", "High", "Server recalculation from submitted events, stricter plausibility validation, signed build/version policy and anomaly reports.", "Developer/Admin"),
        ("Database unavailable or migration fails", "Low/medium", "High", "Connection pool health checks, repeatable migrations, Neon backups/branching and documented restore procedure.", "Developer"),
        ("Mobile form clips controls", "Medium", "Medium", "Dynamic viewport units, safe-area padding, vertical scrolling and tests on short/narrow devices.", "Developer/Tester"),
        ("Third-party model licence is unclear", "Low", "High", "Use CC0 assets where possible, retain attribution file and store source/licence metadata with each asset.", "Developer"),
        ("Backend stack does not meet brief", "High", "High", "Obtain written tutor approval or reimplement the service in Python, Java or PHP before submission.", "Student/Tutor"),
        ("Peer feedback is unavailable before deadline", "Medium", "Medium", "Book named testers early, use the UAT checklist, record device/browser and preserve dated evidence.", "Student"),
    ], [1770, 1000, 900, 4530, 1160])

    h2(doc, "2.7 Use cases")
    add_picture(doc, figs["use_case"], "Figure 1. Use-case model; dashed administration functions are planned.", "Use-case diagram showing guest and player game actions, plus planned administrator functions.")
    h2(doc, "2.8 Data design")
    add_picture(doc, figs["erd"], "Figure 2. Current PostgreSQL entity-relationship model.", "Entity-relationship diagram for players, accounts, sessions, progress, runs, event data, skins and achievements.", width=6.35)
    add_body(doc, "The schema normalises identity, credentials, mutable progression, settings and repeatable run/event records. Foreign keys and cascading deletes protect referential integrity, while checks constrain counts, difficulty values, outcomes and score-related fields. A leaderboard view joins player identity with progress data. The design currently lacks roles, permissions, administrative audit records and feedback entities; these should be added without overloading the existing gameplay event tables.")

    h2(doc, "2.9 Three-tier architecture")
    add_picture(doc, figs["architecture"], "Figure 3. Three-tier deployment and runtime architecture.", "Three-tier diagram showing React and Capacitor clients, Express API, PostgreSQL database and static assets.")
    h2(doc, "2.10 Data-flow diagrams")
    add_picture(doc, figs["dfd0"], "Figure 4. Level 0 context data-flow diagram.", "Context data-flow diagram between players, planned administrators, Math Rush 3D and PostgreSQL.")
    add_picture(doc, figs["dfd1"], "Figure 5. Level 1 data-flow diagram.", "Level 1 data-flow diagram for authentication, local gameplay, run submission, progress and reporting.", width=3.8)
    h2(doc, "2.11 Gameplay flowchart")
    add_picture(doc, figs["flow"], "Figure 6. End-to-end gameplay flow.", "Flowchart from application launch through authentication, levels, gates, obstacles, boss, multiplier lane and result actions.", width=3.4)
    page_break(doc)
    h2(doc, "2.12 Interface wireframes")
    add_picture(doc, figs["wireframes"], "Figure 7. Low-fidelity mobile wireframes for the four principal interface states.", "Four mobile wireframes showing authentication, home and level select, gameplay, and victory results.", width=6.35)
    h2(doc, "2.13 Requirements traceability")
    add_body(doc, "The traceability matrix links each user requirement to the current implementation area and the evidence still needed. It prevents a feature from being treated as complete only because it appears on screen.")
    add_table(doc, ["Requirement", "Implementation area", "Verification", "Current position"], [
        ("UR-01 Authentication", "AuthScreen, authentication API routes, player_accounts and auth_sessions", "Registration, login, invalid input and session-expiry integration tests", "Implemented; formal integration evidence pending"),
        ("UR-02 Progression", "HomeScreen, levelProgress utilities and player_progress", "Unlock-order unit tests and cross-device profile test", "Implemented"),
        ("UR-03 Steering", "useInputController, CrowdRunner and track boundaries", "Touch, mouse and keyboard device matrix", "Implemented; UAT pending"),
        ("UR-04 Arithmetic gates", "mathGates utilities and GateManager", "Expression correctness tests across all difficulties", "Implemented"),
        ("UR-05 Individual casualties", "CrowdRuntime, ObstacleManager and obstacle events", "Collision and crowd-loss regression tests", "Implemented; animation UAT pending"),
        ("UR-06 Boss and finish", "BossManager, boss meter, GameTrack and bonus stages", "Perfect-route balance test and end-state UAT", "Implemented"),
        ("UR-07 Run controls", "Pause overlay, resetGame, replay and next-level actions", "Pause, restart and duplicate-submission checks", "Implemented"),
        ("UR-08 Persistence", "Zustand persistence, player API, run queue and PostgreSQL", "Offline queue plus second-client login test", "Implemented; integration evidence pending"),
        ("UR-09 Rankings", "Leaderboard route, view and statistics modal", "Ordering, filtering and privacy tests", "Partially verified"),
        ("UR-10 Administration", "No protected administrator module identified", "RBAC denial and authorised CRUD tests", "Not implemented"),
    ], [1300, 3440, 2740, 1880])


def add_development(doc: Document) -> None:
    h1(doc, "3. Development technologies and methodology")
    h2(doc, "3.1 Software and tools used")
    add_table(doc, ["Area", "Selected technology", "Purpose and justification"], [
        ("Interface", "React 19 + React DOM", "Component-based menus and overlays; React's state and composition model suits reusable screens and controls (Meta Platforms, 2026)."),
        ("Language", "TypeScript 6", "Static types across browser, server and shared rules reduce mismatched event and API structures (Microsoft, 2026)."),
        ("Build", "Vite 8", "Fast development server and optimised production bundling for a modern React application (Vite, 2026)."),
        ("Styling", "Tailwind CSS 4 + project CSS", "Responsive utility classes support consistent spacing and breakpoints while custom CSS handles game-specific surfaces (Tailwind Labs, 2026)."),
        ("3D", "Three.js + React Three Fiber + Drei", "WebGL rendering with a declarative React scene graph and reusable helpers; R3F is explicitly a React renderer for Three.js (Poimandres, 2026)."),
        ("State", "Zustand", "Small central store with selective subscriptions and persistence; frame-critical transforms remain outside reactive state."),
        ("API", "Node.js + Express 5 + Zod", "REST endpoints, middleware, validation and shared TypeScript. Academically, this selection needs tutor approval because of the brief's language restriction."),
        ("Database", "PostgreSQL 17/18 + pg", "Relational constraints, foreign keys, views, transactions and parameterised SQL; PostgreSQL supports complex queries and transactional integrity (PostgreSQL Global Development Group, 2026)."),
        ("Mobile", "Capacitor 8 + Android/Gradle", "Packages the web-first application in a native Android container and retains access to native platform APIs (Ionic, 2026)."),
        ("Testing", "Vitest + TypeScript + Oxlint", "Fast rule-level regression tests, compile-time checks and static analysis."),
        ("Delivery", "GitHub + GitHub Actions + Render + Neon", "Version history, automated APK workflow, static/API hosting and managed PostgreSQL."),
        ("Design/evidence", "Word, Graphviz and low-fidelity wireframes", "Structured report, repeatable technical diagrams and an initial screen plan; editable Figma evidence remains pending."),
    ], [1450, 2500, 5410])

    h2(doc, "3.2 Alternatives considered")
    add_table(doc, ["Decision", "Selected", "Alternative", "Reasoned comparison"], [
        ("3D integration", "React Three Fiber", "Plain Three.js", "Plain Three.js gives direct control and fewer React concepts, but R3F keeps the scene declarative and componentised while still exposing Three.js objects. The current team skill set and React UI favour R3F."),
        ("State management", "Zustand", "React Context / Redux Toolkit", "Context is sufficient for low-frequency UI state but can become awkward across many game consumers. Redux offers stronger conventions and tooling but adds ceremony. Zustand is compact; careful selectors are still required."),
        ("Server", "Express/TypeScript", "FastAPI/Python, Spring Boot/Java, Laravel/PHP", "Express maximises type sharing and implementation speed. FastAPI would satisfy the brief and provide strong validation/OpenAPI; Spring Boot offers mature structure but higher complexity; Laravel offers rapid CRUD and RBAC patterns. Compliance may outweigh reuse."),
        ("Database access", "Direct pg SQL", "Prisma ORM", "Direct SQL makes schema and queries visible for assessment and avoids ORM overhead. Prisma would improve generated types and migration ergonomics but can hide SQL learning and adds a generation layer."),
        ("Database", "PostgreSQL", "MySQL / SQLite", "PostgreSQL is required by the brief and supports constraints, views and rich querying. SQLite is attractive offline but unsuitable as the shared production leaderboard store."),
        ("Mobile delivery", "Capacitor", "React Native / native Android", "Capacitor reuses the deployed web client and is appropriate for the current project. React Native or native Android could improve device-specific optimisation but would duplicate rendering and interface work."),
        ("Methodology", "Agile iterative", "Waterfall", "The game required repeated balancing and mobile feedback, which benefits from short build-test-review loops. Waterfall offers stable documentation gates but delays gameplay validation until too late."),
    ], [1450, 1800, 2050, 4060])

    h2(doc, "3.3 Frontend implementation")
    add_body(doc, "React components own the screen-level interface, while GameScene creates a level definition and passes it to specialised scene systems. GateManager, ObstacleManager and BossManager respond to the same generated values; CrowdRunner and CrowdRuntime maintain movement and formation state; FollowCamera tracks the action; GameTrack renders the route and finish. React Three Fiber permits these scene responsibilities to remain component-based while the per-frame loop operates on Three.js objects rather than rebuilding ordinary DOM elements (Poimandres, 2026).")
    add_body(doc, "The Zustand store models phase, run stage, difficulty, selected level, profile, settings, rewards and event summaries. Long-lived preferences are persisted locally and reconciled with the authenticated profile. Position and animation data that change every frame are kept in refs and runtime objects to avoid excessive component rerendering. The crowd uses instanced drawing so many bodies can share geometry and material data.")
    add_body(doc, "Responsive behaviour includes touch dragging, mouse/keyboard steering, safe-area insets, capped device pixel ratio, a reduced-effects option and scrollable authentication content. Capacitor provides the native Android shell while preserving the same React application, consistent with its web-first runtime model (Ionic, 2026).")

    h2(doc, "3.4 Backend implementation")
    add_body(doc, "Express exposes health, authentication, player, run and leaderboard routes. Request bodies and parameters are validated with Zod. Helmet configures common security headers, CORS restricts approved client origins, JSON bodies are limited and authentication endpoints are rate-limited. Passwords are hashed with bcrypt at cost 12. This meets OWASP's stated minimum bcrypt work factor of 10, although OWASP prefers Argon2id for new systems where available (OWASP Foundation, 2026a). Random session tokens are returned to the client but only SHA-256 token hashes are stored in the database, limiting direct reuse after a database leak.")
    add_body(doc, "Browser sessions use secure HTTP-only cookies in production, while the Android client uses a bearer token because WebView cookie behaviour differs across origins. Sessions expire after 30 days and can be invalidated. OWASP describes the session identifier as temporarily equivalent to the strongest authentication method, so HTTPS, secure storage, expiry, rotation and logout behaviour must be tested as security controls rather than treated as interface details (OWASP Foundation, 2026b).")
    add_callout(doc, "Compliance decision required", "The assignment brief names PHP, Python or Java for the backend and allows Node.js only as middleware or API glue. The current Express service performs core business and persistence work. Obtain tutor approval in writing or migrate this layer before final submission.", fill="FBEAEA", accent=RED)

    h2(doc, "3.5 Database and SQL")
    add_body(doc, "Six ordered SQL migrations create the current schema. UUID and bigserial keys identify durable entities and events. Enumerated types constrain difficulty, run status and obstacle kinds. Check constraints prevent negative balances and invalid levels; foreign keys enforce ownership; indexes support player history and leaderboard ordering; triggers maintain updated timestamps; and a view exposes leaderboard data without duplicating it. The API uses a PostgreSQL connection pool capped at ten clients and parameterised queries rather than interpolating user input.")
    add_body(doc, "The current model satisfies substantial persistence needs, but it is not yet the complete database described by the assignment. A future migration should add roles, permissions or a role column, administrator audit actions, player feedback, content-management entities where appropriate and report-friendly indexes. Run events should remain gameplay analytics; administrative actions require a separate immutable audit trail with actor, action, target, timestamp and before/after metadata.")

    h2(doc, "3.6 Development methodology")
    add_body(doc, "An Agile iterative lifecycle is the best fit because movement feel, crowd behaviour, arithmetic readability and level difficulty cannot be validated reliably from documentation alone. The Agile Manifesto prioritises working software, collaboration and responding to change while retaining value in processes and plans (Beck et al., 2001). The project uses those principles without claiming a complete formal Scrum implementation.")
    add_number(doc, "Discover: interpret the assignment, inspect comparable crowd runners, identify the learning problem and record constraints.")
    add_number(doc, "Design: write requirements, model data and flow, plan mobile wireframes, define level-generation and reward rules.")
    add_number(doc, "Build a vertical slice: home screen, one playable level, gates, crowd control, boss and result.")
    add_number(doc, "Expand: add database accounts, progression, skins, leaderboard, difficulty bands, obstacles and Android packaging.")
    add_number(doc, "Verify each iteration: automated rule tests, type checks, build checks, browser/device testing and database/API checks.")
    add_number(doc, "Review and adapt: interpret peer feedback, prioritise defects, update the SDD and preserve evidence in GitHub.")
    add_body(doc, "A practical two-week iteration can use a prioritised backlog, a defined acceptance condition, a short implementation period, automated checks, a device demonstration and a retrospective. Scrum's emphasis on inspection and adaptation supports this cadence, but roles and events should only be claimed if they were actually performed (Schwaber and Sutherland, 2020).")

    h2(doc, "3.7 Version control and delivery")
    add_body(doc, "Git tracks source changes and GitHub hosts the MathRush3D repository. Commits should describe one coherent change, and feature branches or pull requests should be used for risky work. Tags should identify submitted versions. Secrets such as DATABASE_URL must remain in deployment environment variables and must not be committed. GitHub Actions can automate build and test workflows and retain APK artifacts; GitHub describes Actions as a way to execute customised software-development workflows, including CI/CD (GitHub, 2026).")
    add_bullet(doc, "Minimum evidence for the final appendix: repository URL, commit history screenshot, branch/pull-request example, passing test/build workflow and tagged release.")
    add_bullet(doc, "The Android pipeline must use a supported Node version, Java 21, Android SDK 36 and the repository's Gradle wrapper.")
    add_bullet(doc, "Render environment variables must separately configure DATABASE_URL, database SSL, NODE_ENV and the exact comma-separated client origins.")
    h2(doc, "3.8 Data protection, ethics and asset licensing")
    add_body(doc, "The service should collect only the data required for authentication, progression and competitive results. Email addresses and password hashes must never be exposed through leaderboard responses. Session records should expire and stale sessions should be removable. The final system should document a retention period for inactive accounts and provide an account-deletion process that relies on the schema's cascading relationships.")
    add_body(doc, "The game does not currently include chat, advertising, payments or location tracking, which reduces safeguarding and privacy risk. Display names still require validation because they are visible to other players. Arithmetic difficulty should be presented as a game setting rather than a judgement of ability, and feedback should avoid humiliating language when a player loses.")
    add_body(doc, "External models and textures must retain source and licence records. The current project uses local attribution metadata for Kenney and Quaternius assets. Any replacement asset should be accepted only after its redistribution and commercial-use terms are verified.")


def add_testing(doc: Document) -> None:
    h1(doc, "4. Testing and user acceptance")
    h2(doc, "4.1 Testing strategy")
    add_body(doc, "Testing is layered so that fast deterministic checks catch rule regressions before slower device and user testing. Unit tests cover balance, crowd formation, level progression, speed progression and boss-meter results. TypeScript checks validate browser and server types. API integration tests should create a temporary account, authenticate, submit a run, retrieve progress and verify leaderboard ordering against a disposable database. End-to-end tests should then exercise registration, a complete level, pause/restart, next level and logout through the rendered interface.")
    h2(doc, "4.2 Current automated evidence")
    add_table(doc, ["Check", "Command", "Observed result", "Date"], [
        ("Gameplay utility tests", "npm test", "5 test files passed; 30 of 30 tests passed in Vitest.", "4 Aug 2026"),
        ("Browser and server type checks", "npm run typecheck", "Completed successfully as part of the production build.", "4 Aug 2026"),
        ("Production build", "npm run build", "Completed successfully. Vite reported a 1,275.81 kB JavaScript bundle before gzip and issued a large-chunk warning.", "4 Aug 2026"),
        ("Database/API integration", "Planned test script", "No automated integration suite was identified in the inspected repository.", "Pending"),
        ("Android device matrix", "Manual installation/UAT", "Informal testing occurred, but a formal device matrix and signed evidence are still required.", "Pending"),
    ], [1900, 1900, 3900, 1660])

    h2(doc, "4.3 Test data and environments")
    for item in (
        "Desktop web: current Chrome and Edge at 1366x768 and 1920x1080.",
        "Mobile web: 360x640, 390x844 and 412x915 responsive viewports, including on-screen keyboard behaviour.",
        "Android: at least one low/mid-range device and one newer device, recording OS version, memory and screen resolution.",
        "API/database: local PostgreSQL container and production Neon database with separate non-production test accounts.",
        "Adverse network: offline start, delayed first request, temporary API failure and retry after reconnection.",
    ):
        add_bullet(doc, item)

    h2(doc, "4.4 Peer tester plan")
    add_body(doc, "At least two peers should test independently. One tester should focus on first-time usability and mathematical clarity; the other should focus on mobile controls, difficulty and defects. The student should not coach testers beyond the instructions visible in the product. Record the tester's consent, device, browser/app version, date, observations, severity and suggested improvement. Do not replace observed evidence with assumed feedback.")
    add_table(doc, ["Tester", "Profile", "Device/environment", "Session", "Evidence"], [
        ("[Tester 1 name]", "First-time player / learner", "[Device, OS, browser]", "Register, complete Easy 1-2, inspect result and leaderboard", "Pending"),
        ("[Tester 2 name]", "Mobile/gameplay tester", "[Device, OS, APK version]", "Test squeeze control, obstacles, pause/restart and boss/bonus sequence", "Pending"),
    ], [1600, 1800, 2050, 2600, 1310])

    h2(doc, "4.5 User acceptance checklist")
    uat_rows = [
        ("UAT-01", "Register with valid name, email and password.", "Account is created and home screen loads.", "Pending"),
        ("UAT-02", "Enter invalid or duplicate credentials.", "Specific, readable validation/error response; no account corruption.", "Pending"),
        ("UAT-03", "Reload and log in on another supported client.", "Saved progression, coins, skin and statistics are restored.", "Pending"),
        ("UAT-04", "Select an unlocked and a locked level.", "Unlocked level starts; locked level cannot start and explains prerequisite.", "Pending"),
        ("UAT-05", "Play using touch drag and keyboard.", "Crowd follows input, remains in bounds and page does not scroll during play.", "Pending"),
        ("UAT-06", "Choose gates across all difficulties.", "Expressions are legible, integer-valued and update the crowd correctly.", "Pending"),
        ("UAT-07", "Press fully against each track wall near hazards.", "Crowd visibly compresses and skilled positioning reduces casualties.", "Pending"),
        ("UAT-08", "Collide with each obstacle type and an enemy crowd.", "Correct individuals/damage are removed once; animation and sound match the event.", "Pending"),
        ("UAT-09", "Reach the boss with an optimal route.", "Level is beatable; timing meter is readable and boss damage is understandable.", "Pending"),
        ("UAT-10", "Complete the bonus multiplier lane.", "Crowd is consumed stage by stage, the track ends visibly and rewards match the result.", "Pending"),
        ("UAT-11", "Pause, restart, replay, next level and home.", "Each action changes state once and does not duplicate a run submission.", "Pending"),
        ("UAT-12", "Disable sound/vibration and enable reduced effects.", "Settings apply immediately and persist after restart/login.", "Pending"),
        ("UAT-13", "Disconnect the network, finish a run, then reconnect.", "Game remains playable; run queues and later synchronises once.", "Pending"),
        ("UAT-14", "Use leaderboard difficulty filter/search function.", "Results are correct and inputs cannot expose another player's private data.", "Pending / feature gap"),
        ("UAT-15", "Attempt administrator functions as player/admin.", "Player is denied; authorised administrator can use audited CRUD.", "Blocked: not implemented"),
    ]
    add_table(doc, ["ID", "Test action", "Expected result", "Result"], uat_rows, [850, 3150, 3990, 1370])

    h2(doc, "4.6 Performance evaluation plan")
    add_body(doc, "Performance claims should be supported by repeatable measurements rather than visual judgement. Each test must record the application version, device, operating system, browser or APK version, selected level, crowd size and reduced-effects setting.")
    add_table(doc, ["Metric", "Target", "Method", "Test point", "Action if missed"], [
        ("Frame rate", "At least 45 FPS on the selected mid-range Android device", "Record a 60-second run with browser/Android profiling tools", "Largest crowd, gate transition and boss battle", "Reduce effects, profile allocations and review draw calls"),
        ("Frame time", "95th percentile below 22 ms", "Capture performance trace and inspect long frames", "Gate application and crowd casualties", "Remove synchronous work and repeated object creation"),
        ("Initial load", "Playable screen within 5 seconds on normal 4G after service wake-up", "Cold-cache timed test", "First visit and first APK launch", "Compress assets, split bundles and improve loading feedback"),
        ("API latency", "Warm health/profile request below 500 ms", "Twenty production requests with median and 95th percentile", "Login and post-run synchronisation", "Inspect database queries, indexes and hosting region"),
        ("Memory stability", "No continuing growth after three complete runs", "Compare memory snapshots before and after repeated runs", "Replay and next-level sequence", "Dispose scene resources and remove stale listeners"),
        ("Offline recovery", "One eventual run record with no duplicates", "Finish offline, reconnect and query run history", "Queued submission", "Fix idempotency and retry handling"),
    ], [1550, 2050, 2100, 1900, 1760])

    h2(doc, "4.7 Defect and feedback handling")
    add_body(doc, "Each failed check should become a backlog item with reproduction steps, expected and actual behaviour, device/build information, severity and supporting image or video. Feedback should be interpreted rather than copied directly into a change: repeated observations, task failure and high-severity defects outrank individual visual preferences. After a fix, the original check and relevant regression tests must be repeated. The final report should show at least one traceable chain from peer observation to decision, implementation, retest and evaluation.")


def add_compliance(doc: Document) -> None:
    h1(doc, "5. Current assignment coverage and gaps")
    add_body(doc, "This matrix compares the inspected project and this report with the supplied brief. 'Partial' means that some evidence or functionality exists but the criterion should not yet be claimed as achieved.")
    add_table(doc, ["Criterion", "Current evidence", "Assessment", "Next evidence/action"], [
        ("P1 Problem and requirements", "Problem definition, stakeholders, user/system requirements and diagrams in this report.", "Substantial draft", "Validate requirements with named users and update after feedback."),
        ("P2 Risks", "Risk register covers gameplay, performance, security, data, licensing, mobile and compliance.", "Substantial draft", "Add owner review dates and evidence of mitigations."),
        ("M1 SDD analysis", "SDD connects the learning problem to architecture, data and game flow.", "Partial", "Add requirement traceability to implementation screenshots/commits."),
        ("P3 Tools and techniques", "Selected stack, purpose and official research references documented.", "Substantial draft", "Add development screenshots and actual usage examples."),
        ("M2 Justification", "Alternatives compared across 3D, state, backend, data, mobile and methodology.", "Partial", "Deepen comparison using measured build/performance evidence."),
        ("D1 Evaluation of solution/methodology", "Only preliminary evaluation is possible at this stage.", "Pending", "Evaluate after UAT and compare outcomes against alternatives."),
        ("P4 Formal peer presentation", "Tester plan exists; no presentation evidence supplied.", "Pending", "Deliver presentation and attach date, audience, slides and feedback."),
        ("P5 Functional app/support docs", "Playable web/APK, README, database notes and this revised report exist.", "Partial", "Add user/admin guide, screenshots and backend compliance resolution."),
        ("M3 Interpret feedback", "Feedback method is defined but no completed peer results exist.", "Pending", "Complete UAT, group themes and justify accepted/rejected changes."),
        ("M4 Develop against SDD", "Modular implementation and passing rule/type checks exist.", "Partial", "Map requirements to commits, tests and screenshots."),
        ("P6 Review performance", "Performance design choices are described.", "Pending/partial", "Capture FPS, load time, bundle size and device results."),
        ("M5 Critical review of stages/risks", "Initial risks and lifecycle are documented.", "Pending/partial", "Compare estimated and actual risks after remaining iterations."),
        ("D2 Improvements and future work", "Open gaps are prioritised below but not yet feedback-validated.", "Pending", "Link final recommendations to test and peer evidence."),
    ], [1150, 3620, 1400, 3190])

    h2(doc, "5.1 Mandatory feature gap analysis")
    add_table(doc, ["Brief requirement", "Current state", "Required action"], [
        ("PostgreSQL and SQL", "Implemented through migrations, parameterised queries, constraints and a leaderboard view.", "Add schema/query evidence and backup/restore test."),
        ("Frontend React + CSS framework", "Implemented with React, Tailwind CSS and custom CSS.", "Add responsive screenshots and accessibility evidence."),
        ("Backend PHP/Python/Java", "Core backend currently uses Node.js/Express/TypeScript.", "Obtain explicit approval or migrate to FastAPI, Spring Boot or PHP before submission."),
        ("CRUD", "Account/profile, progress, skins and run creation/read paths exist; complete admin CRUD is absent.", "Define core entities and implement protected create/read/update/delete workflows."),
        ("RBAC", "No role/permission model or admin interface identified.", "Add roles, authorisation middleware, admin UI and denial tests."),
        ("Validation/error handling", "Zod, database constraints and central API error handling exist.", "Add integration tests and user-friendly error mapping."),
        ("Search/filter", "Leaderboard supports limited retrieval/filter behaviour; general search is incomplete.", "Add documented server-side filters/search with pagination."),
        ("Dashboard/report", "Player statistics and leaderboard exist.", "Add administrator reporting and exportable summary if required."),
        ("Responsive UI", "Mobile styles, safe areas, scroll fixes and touch controls exist.", "Complete formal viewport/device matrix."),
        ("Audit logs", "Gate and obstacle events are analytics, not administrator audit records.", "Create separate audit_log entity and record protected changes."),
        ("Feedback", "No persistent feedback entity or completed peer evidence identified.", "Add feedback table/form if required and complete UAT records."),
    ], [2450, 3590, 3320])

    h2(doc, "5.2 Prioritised next stage")
    add_number(doc, "Resolve the backend-language compliance decision with the tutor before investing further in the existing service.")
    add_number(doc, "Implement RBAC, administrator authentication/authorisation, protected CRUD, audit logs and denial tests.")
    add_number(doc, "Add API/database integration tests, server-side reward plausibility checks and repeatable test data.")
    add_number(doc, "Complete search/filter, administrator reports and a persistent feedback workflow where the full guideline requires them.")
    add_number(doc, "Run peer UAT on web and Android, interpret results, implement justified changes and retest.")
    add_number(doc, "Capture final evidence: diagrams, SQL, commits, CI, responsive screens, performance metrics and versioned deployment URLs.")


def add_conclusion_and_refs(doc: Document) -> None:
    h1(doc, "6. Conclusion", new_page=False)
    add_body(doc, "Math Rush 3D already demonstrates a coherent, technically ambitious application: a React-based 3D game, a persistent PostgreSQL data model, secure account handling, responsive Android delivery and automated rule checks. Its strongest design decision is the use of a single generated level definition and shared reward rules to reduce disagreement between gameplay systems and stored outcomes. The clearest weaknesses are academic compliance around the backend language, missing role-based administration and the absence of completed peer/evaluation evidence.")
    add_body(doc, "This report covers the topics listed in the supplied project guideline: executive summary, SDD, problem and stakeholder analysis, requirements, measurable success criteria, risk review, use-case, ER, architecture, DFD, flowchart and wireframe designs, requirement traceability, software choices and alternatives, frontend, backend and database discussion, methodology, version control, ethics, performance evaluation and a peer UAT plan. The report does not claim completion of evidence that has not been supplied. Peer testing, measured device performance, administration functions and the backend-language decision remain necessary before the related higher-grade criteria can be evidenced confidently.")

    h1(doc, "References")
    references = [
        ("Beck, K. et al. (2001) Manifesto for Agile Software Development. Available at: ", "https://agilemanifesto.org/", "Agile Manifesto", " (Accessed: 21 July 2026)."),
        ("GitHub (2026) GitHub Actions documentation. Available at: ", "https://docs.github.com/en/actions", "GitHub Actions documentation", " (Accessed: 21 July 2026)."),
        ("Ionic (2026) Capacitor: Cross-platform native runtime for web apps. Available at: ", "https://capacitorjs.com/docs", "Capacitor documentation", " (Accessed: 21 July 2026)."),
        ("Meta Platforms (2026) React Quick Start. Available at: ", "https://react.dev/learn", "React Quick Start", " (Accessed: 21 July 2026)."),
        ("Microsoft (2026) TypeScript documentation. Available at: ", "https://www.typescriptlang.org/docs/", "TypeScript documentation", " (Accessed: 21 July 2026)."),
        ("OpenJS Foundation (2026) Express 5: Installing and TypeScript. Available at: ", "https://expressjs.com/en/starter/installing/", "Express documentation", " (Accessed: 21 July 2026)."),
        ("OWASP Foundation (2026a) Password Storage Cheat Sheet. Available at: ", "https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html", "OWASP Password Storage Cheat Sheet", " (Accessed: 21 July 2026)."),
        ("OWASP Foundation (2026b) Session Management Cheat Sheet. Available at: ", "https://cheatsheetseries.owasp.org/cheatsheets/Session_Management_Cheat_Sheet.html", "OWASP Session Management Cheat Sheet", " (Accessed: 21 July 2026)."),
        ("Poimandres (2026) React Three Fiber: Introduction. Available at: ", "https://r3f.docs.pmnd.rs/getting-started/introduction", "React Three Fiber introduction", " (Accessed: 21 July 2026)."),
        ("PostgreSQL Global Development Group (2026) What is PostgreSQL? PostgreSQL 18 Documentation. Available at: ", "https://www.postgresql.org/docs/current/intro-whatis.html", "PostgreSQL documentation", " (Accessed: 21 July 2026)."),
        ("Schwaber, K. and Sutherland, J. (2020) The Scrum Guide. Available at: ", "https://scrumguides.org/scrum-guide.html", "The Scrum Guide", " (Accessed: 21 July 2026)."),
        ("Tailwind Labs (2026) Styling with utility classes. Available at: ", "https://tailwindcss.com/docs/styling-with-utility-classes", "Tailwind CSS documentation", " (Accessed: 21 July 2026)."),
        ("Three.js Authors (2026) Three.js manual. Available at: ", "https://threejs.org/manual/", "Three.js manual", " (Accessed: 21 July 2026)."),
        ("Vitest (2026) Getting Started. Available at: ", "https://vitest.dev/guide/", "Vitest guide", " (Accessed: 21 July 2026)."),
        ("Vite (2026) Getting Started. Available at: ", "https://vite.dev/guide/", "Vite guide", " (Accessed: 21 July 2026)."),
    ]
    for prefix, url, link_text, suffix in references:
        p = doc.add_paragraph(style="Reference Entry")
        r = p.add_run(prefix)
        set_run_font(r, size=11, color=INK)
        add_hyperlink(p, link_text, url)
        r = p.add_run(suffix)
        set_run_font(r, size=11, color=INK)

    h1(doc, "Appendix A. Evidence still to insert")
    for item in (
        "The remaining tutor guideline and any revised section order.",
        "Figma project link or exported annotated wireframes if specifically required.",
        "Dated screenshots of registration, home/level select, gameplay, boss meter, bonus finish, results, skins, statistics and leaderboard.",
        "GitHub repository, commit history, pull request/branch and passing CI evidence.",
        "Database migration and representative SQL query evidence with sensitive values redacted.",
        "Named peer presentation and UAT evidence, feedback interpretation, implemented changes and retest results.",
        "Measured performance results on the specified desktop and Android device matrix.",
        "Written tutor decision on the Node.js/Express backend or evidence of migration to an accepted language.",
    ):
        add_bullet(doc, item)

    h1(doc, "Appendix B. Project technology snapshot")
    add_table(doc, ["Layer", "Current implementation"], [
        ("Web client", "React 19.2, React DOM, TypeScript 6, Vite 8, Tailwind CSS 4, HTML/CSS"),
        ("3D/game", "Three.js 0.185, React Three Fiber 9.6, Drei 10.7, WebGL, GLB/GLTF, GPU instancing"),
        ("State/audio", "Zustand 5 with persistence; Web Audio API procedural effects"),
        ("API/security", "Node.js, Express 5.2, Zod 4, Helmet, CORS, rate limiting, bcryptjs, crypto session tokens"),
        ("Data", "PostgreSQL, pg connection pool, six SQL migrations, Neon production, Docker local development"),
        ("Android", "Capacitor 8.4, Java/AndroidX, Gradle, min SDK 24, target/compile SDK 36, immersive fullscreen"),
        ("Quality", "Vitest, TypeScript, Oxlint, Git/GitHub, GitHub Actions"),
        ("Hosting", "Render static web client, Render API service, Neon PostgreSQL"),
        ("Assets", "CC0 Kenney and Quaternius GLB assets with local attribution metadata"),
    ], [2100, 7260])


def build() -> None:
    WORK.mkdir(parents=True, exist_ok=True)
    FIGURES.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    figs = make_diagrams()
    doc = Document()
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_executive_summary(doc)
    add_intro(doc)
    add_sdd(doc, figs)
    add_development(doc)
    add_testing(doc)
    add_compliance(doc)
    add_conclusion_and_refs(doc)

    doc.core_properties.title = "Math Rush 3D Application Development Report"
    doc.core_properties.subject = "Unit 22 Application Development report"
    doc.core_properties.author = "Vincent Escandallo Castillo"
    doc.core_properties.keywords = "Math Rush 3D, application development, SDD, React, PostgreSQL"
    doc.core_properties.comments = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
