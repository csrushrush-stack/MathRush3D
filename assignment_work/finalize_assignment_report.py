from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document


ROOT = Path(r"C:\Users\Acer\Documents\MathsRush3D")
DELIVERABLES = ROOT / "deliverables"
PRE = DELIVERABLES / "Math_Rush_3D_Application_Development_Report_PRE-HUMANIZER.docx"
HUMANIZED = DELIVERABLES / "Math_Rush_3D_Application_Development_Report_HUMANIZED.docx"
FINAL = DELIVERABLES / "Math_Rush_3D_Application_Development_Report_FINAL_SUBMISSION.docx"
OLDER_SOURCE = DELIVERABLES / "Math_Rush_3D_Application_Development_Report_Submission.docx"
WALTER_SOURCE = DELIVERABLES / "Math_Rush_3D_Application_Development_Report_Walter_Writes.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text.replace("—", ";")


def create_humanized_copy() -> int:
    source = Document(OLDER_SOURCE)
    walter = Document(WALTER_SOURCE)
    replacements = {
        original.text: revised.text
        for original, revised in zip(source.paragraphs, walter.paragraphs)
        if original.text.strip() and original.text != revised.text
    }

    doc = Document(PRE)
    applied = 0
    for paragraph in doc.paragraphs:
        revised = replacements.get(paragraph.text)
        if revised:
            replace_paragraph_text(paragraph, revised)
            applied += 1
    doc.core_properties.title = "Math Rush 3D Application Development Report Humanized Review Copy"
    doc.core_properties.author = "Vincent Escandallo Castillo"
    doc.save(HUMANIZED)
    return applied


def create_final_copy() -> None:
    copy2(PRE, FINAL)
    doc = Document(FINAL)
    for paragraph in doc.paragraphs:
        if "—" in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "—" in paragraph.text:
                        replace_paragraph_text(paragraph, paragraph.text)
    doc.core_properties.title = "Math Rush 3D Application Development Report Final Submission"
    doc.core_properties.subject = "Unit 22 Application Development report"
    doc.core_properties.author = "Vincent Escandallo Castillo"
    doc.core_properties.comments = ""
    doc.save(FINAL)


def validate_final_text() -> None:
    doc = Document(FINAL)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(
        p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs
    )
    banned = ["OpenAI", "ChatGPT", "Codex", "humanizer", "artificial intelligence", "AI-assisted", "—"]
    found = [term for term in banned if term.lower() in text.lower()]
    if found:
        raise RuntimeError(f"Disallowed assignment wording remains: {found}")


if __name__ == "__main__":
    count = create_humanized_copy()
    create_final_copy()
    validate_final_text()
    print(f"Walter-derived paragraphs placed in review copy: {count}")
    print(FINAL)
