import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\deliverables"
    r"\Math_Rush_3D_Application_Development_Report_Submission.docx"
)
RESULTS = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\assignment_work"
    r"\walter_results.json"
)
OUTPUT = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\deliverables"
    r"\Math_Rush_3D_Application_Development_Report_Walter_Writes.docx"
)
AUDIT = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\assignment_work"
    r"\walter_report_audit.json"
)


def words(text):
    return len(text.split())


def replace_text_preserving_first_run_format(paragraph, text):
    first_properties = None
    if paragraph.runs and paragraph.runs[0]._element.rPr is not None:
        first_properties = deepcopy(paragraph.runs[0]._element.rPr)

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    run = paragraph.add_run(text)
    if first_properties is not None:
        run._element.insert(0, first_properties)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main():
    document = Document(SOURCE)
    results = json.loads(RESULTS.read_text(encoding="utf-8"))
    audit = []

    for result in results:
        index = int(result["paragraph_index"])
        source = result["source"].strip()
        returned = result.get("returned", "").strip()
        source_words = max(words(source), 1)
        returned_words = words(returned)
        ratio = returned_words / source_words

        # Walter Lite sometimes returns only an opening fragment or repeats stale
        # content. Keep the original paragraph when the result is not plausibly
        # complete rather than inserting damaged text.
        accepted = (
            bool(returned)
            and 0.70 <= ratio <= 2.00
            and "Variant 1" not in returned
            and len(returned) >= 20
        )

        if accepted:
            replace_text_preserving_first_run_format(document.paragraphs[index], returned)

        audit.append(
            {
                "paragraph_index": index,
                "source_words": source_words,
                "returned_words": returned_words,
                "ratio": round(ratio, 3),
                "accepted": accepted,
            }
        )

    # Reassert the requested typography in every table without changing content.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        rpr = run._element.get_or_add_rPr()
                        rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
                        rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                        rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    document.save(OUTPUT)
    AUDIT.write_text(json.dumps(audit, indent=2), encoding="utf-8")
    print(f"Saved {OUTPUT}")
    print(f"Accepted {sum(item['accepted'] for item in audit)} of {len(audit)} results")
    print(
        "Retained original paragraphs:",
        [item["paragraph_index"] for item in audit if not item["accepted"]],
    )


if __name__ == "__main__":
    main()
