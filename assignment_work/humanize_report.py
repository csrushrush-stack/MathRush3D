from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\deliverables"
    r"\Math_Rush_3D_Application_Development_Report_Submission.docx"
)
OUTPUT = Path(
    r"C:\Users\Acer\Documents\MathsRush3D\deliverables"
    r"\Math_Rush_3D_Application_Development_Report_Humanized.docx"
)


REWRITES = {
    19: (
        "Draft status: This report covers the sections listed in the two-page project "
        "guideline and the attached assignment brief. Some final evidence is still "
        "outstanding, so the report should not yet be treated as complete. Before "
        "submission, the student details, peer-testing evidence, final evaluation and "
        "any later guidance from the tutor must be added."
    ),
    37: (
        "Math Rush 3D is a mobile-first 3D crowd runner that combines quick arcade "
        "gameplay with short arithmetic decisions. I chose this idea because basic "
        "arithmetic practice can feel repetitive when the answer has no visible effect. "
        "Here, addition, subtraction, multiplication and division appear directly on "
        "the route. The gate selected by the player changes the crowd immediately, and "
        "that crowd must then survive the obstacles, defeat the boss and reach the best "
        "possible multiplier."
    ),
    38: (
        "The game is available as a responsive website and as an Android application. "
        "The interface uses React, TypeScript and Tailwind CSS, while Three.js and React "
        "Three Fiber handle the 3D scene. Zustand stores the main interface and run "
        "state, but fast-changing animation values stay outside React state to avoid "
        "unnecessary updates. The server is an Express REST API written in TypeScript. "
        "It uses Zod validation, rate limiting, security headers and session-based "
        "authentication. PostgreSQL stores accounts, settings, progress, run results, "
        "gate choices, obstacle events, skins, achievements and leaderboard data. "
        "Render and Neon host the production version, while Docker can be used locally."
    ),
    39: (
        "A player first registers or logs in, then chooses a difficulty and an unlocked "
        "level. During a run, the player guides a grouped crowd through ten pairs of "
        "arithmetic gates and a set of obstacles designed for that difficulty. Hazards "
        "and enemy crowds remove individual runners. Just before the boss, a timing "
        "meter gives the player a chance to earn an extra crowd boost. The boss fight "
        "then uses the level's balance rules to reduce the crowd. If the player wins, "
        "the remaining runners enter a multiplier lane and are spent stage by stage for "
        "bonus points. The result screen shows the score, coins, stars and progression, "
        "with an option to continue straight to the next level."
    ),
    40: (
        "The current version already has a strong working base. It includes four "
        "difficulty groups with five levels in each, responsive controls, GPU-instanced "
        "crowds, movement-plane collision checks, balanced boss calculations, skins, "
        "persistent accounts and an online leaderboard. On 21 July 2026, all 30 "
        "automated tests passed, together with the browser and server TypeScript checks. "
        "These results show that the balance, formation, progression, speed and boss "
        "meter utilities are internally consistent. They do not, however, replace "
        "testing on real devices, accessibility checks, security testing or user "
        "acceptance testing."
    ),
    41: (
        "A few important assignment requirements are still open. The brief expects PHP, "
        "Python or Java for the backend and only mentions Node.js as middleware or API "
        "glue. Because the present Express server handles the main business and database "
        "logic, written approval from the tutor is needed unless the backend is moved to "
        "an accepted language. The project also lacks role-based access control, an "
        "administrator interface, complete CRUD for the main entities, a formal audit "
        "trail, stored feedback and finished peer-review evidence. These gaps are stated "
        "clearly so that unfinished work is not presented as complete."
    ),
    42: (
        "The next sensible step is to settle the backend-language question with the "
        "tutor. After that, development should focus on the administrator role and "
        "management screens, the remaining CRUD, search and reporting functions, and "
        "peer testing using the UAT checklist. Dated screenshots and clear GitHub "
        "evidence should also be collected. This report gives those later tasks a "
        "traceable starting point and can be updated when the rest of the guideline is "
        "provided."
    ),
    45: (
        "The assignment presents the project as an entry for a GameForge Studios "
        "hackathon. The aim is to build a browser game that addresses a real problem "
        "while also showing sound application-development practice, relational database "
        "use and honest evaluation. Math Rush 3D focuses on the difficulty of keeping "
        "learners interested in repeated mental-arithmetic practice, especially when "
        "ordinary exercises offer little choice or immediate visual feedback."
    ),
    47: (
        "The aim is to produce a responsive and replayable arithmetic game in which "
        "success depends on both mathematical choices and movement skill. A secure data "
        "service should also keep progress, settings and competitive results consistent "
        "between the web and Android versions."
    ),
    58: (
        "The project currently covers a single-player crowd runner, arithmetic gates, a "
        "local 3D simulation, account-based saving and competitive rankings. Real-time "
        "multiplayer, payments, advertisements, social messaging, user-created levels "
        "and an iOS build are outside the present scope. The brief also expects an "
        "administrator role and a formal feedback process. These remain planned because "
        "neither was found in the version of the repository reviewed for this report."
    ),
    61: (
        "Worksheets are useful for checking arithmetic answers, but they do not always "
        "encourage learners to practise voluntarily. They also separate the calculation "
        "from any immediate consequence. Math Rush 3D makes the calculation part of a "
        "spatial choice: the player reads two expressions, moves towards one and sees "
        "the crowd change at once. This only works if the game remains fair. Symbols "
        "must be clear on small screens, generated levels must stay beatable, and good "
        "steering should reduce losses without making every obstacle harmless."
    ),
    85: (
        "The core game can run locally when the API is unavailable, although an internet "
        "connection is required to synchronise an account. Render's free hosting may "
        "need time to wake up, so the first request can be noticeably slower. Mobile GPU "
        "and memory limits also affect the 3D scene, which is why instancing and bounded "
        "visual effects are more useful here than highly detailed models. From an "
        "academic point of view, the main constraint is the Express backend. Its "
        "acceptability under the brief must be confirmed with the tutor before the final "
        "submission."
    ),
    94: (
        "The database separates identity, login details, changeable progress, settings "
        "and repeatable run or event records. Foreign keys and cascading deletes keep "
        "related data consistent, while check constraints reject invalid counts, levels, "
        "difficulties, outcomes and score values. A database view combines player and "
        "progress data for the leaderboard. Roles, permissions, administrator audit "
        "records and feedback are still missing. They should be introduced as separate "
        "entities instead of being mixed into the existing gameplay-event tables."
    ),
    110: (
        "The traceability matrix connects each user requirement to the relevant part of "
        "the implementation and identifies the evidence that is still missing. This "
        "makes it harder to mark a feature as finished simply because it appears on the "
        "screen."
    ),
    118: (
        "React components control the main screens. When a level begins, GameScene builds "
        "the level definition and supplies it to the specialised scene systems. "
        "GateManager, ObstacleManager and BossManager therefore work from the same set of "
        "generated values. CrowdRunner and CrowdRuntime handle movement and formation, "
        "FollowCamera tracks the run, and GameTrack draws the route and finish area. "
        "React Three Fiber keeps these responsibilities component-based while allowing "
        "the frame loop to update Three.js objects directly instead of rebuilding normal "
        "DOM content (Poimandres, 2026)."
    ),
    119: (
        "The Zustand store keeps the current phase, run stage, difficulty, selected "
        "level, profile, settings, rewards and event summaries. Preferences that need to "
        "survive a restart are saved locally and later reconciled with the signed-in "
        "profile. Positions and animation values change too often to belong in reactive "
        "state, so they remain in refs and runtime objects. The crowd is rendered with "
        "instancing, allowing many runners to share the same geometry and material."
    ),
    120: (
        "The responsive design supports touch dragging, mouse and keyboard steering, "
        "safe-area insets, a limited device pixel ratio, reduced effects and scrolling "
        "within the authentication screen. Capacitor wraps the same React application "
        "for Android, so the project can keep a web-first codebase while still producing "
        "a native package (Ionic, 2026)."
    ),
    122: (
        "The Express API provides health, authentication, player, run and leaderboard "
        "routes. Zod checks request bodies and parameters before they reach the main "
        "logic. Helmet adds common security headers, CORS limits approved client origins, "
        "JSON request size is restricted and authentication requests are rate-limited. "
        "Passwords use bcrypt with a cost of 12, which is above OWASP's stated minimum "
        "work factor of 10, although Argon2id is preferred for new systems when it is "
        "available (OWASP Foundation, 2026a). The client receives a random session token, "
        "but the database stores only its SHA-256 hash. This reduces the value of stolen "
        "session records."
    ),
    123: (
        "Production browser sessions use secure HTTP-only cookies. The Android client "
        "uses a bearer token because cookie behaviour can differ inside a WebView and "
        "across origins. A session lasts for up to 30 days and can be invalidated. Since "
        "OWASP treats a session identifier as temporarily equivalent to the user's "
        "authentication, HTTPS, secure storage, expiry, rotation and logout must all be "
        "tested as security controls, not treated as minor interface details (OWASP "
        "Foundation, 2026b)."
    ),
    124: (
        "Compliance decision required: the brief names PHP, Python or Java as backend "
        "languages and allows Node.js only for middleware or API glue. The current "
        "Express service performs the main business and persistence work. Written tutor "
        "approval is therefore required, otherwise this layer should be migrated before "
        "the final submission."
    ),
    126: (
        "The current schema is created through six ordered SQL migrations. UUID and "
        "bigserial keys identify long-lived records and event rows. Enumerated types "
        "limit difficulty, run status and obstacle values. Check constraints block "
        "negative balances and invalid levels, foreign keys enforce ownership, indexes "
        "support run history and leaderboard ordering, and triggers maintain update "
        "timestamps. A view exposes the leaderboard without storing the same information "
        "twice. The API uses parameterised queries and a PostgreSQL pool limited to ten "
        "connections."
    ),
    127: (
        "This model already covers most gameplay persistence, but it does not yet meet "
        "every database requirement in the assignment. A later migration should add "
        "roles or permissions, administrator audit actions, player feedback, suitable "
        "content-management entities and indexes for reports. Gameplay events should "
        "remain separate from administrator activity. The audit trail needs to record "
        "the actor, action, target, time and before-and-after information in an immutable "
        "form."
    ),
    129: (
        "An iterative Agile approach suits this project because movement, crowd behaviour, "
        "gate readability and level difficulty cannot be judged properly from plans "
        "alone. They need to be built, played and adjusted. The Agile Manifesto values "
        "working software, collaboration and responding to change while still "
        "recognising the value of processes and plans (Beck et al., 2001). The project "
        "follows those ideas, but it does not claim that a complete formal Scrum process "
        "was used."
    ),
    136: (
        "A practical two-week cycle could begin with a prioritised backlog and a clear "
        "acceptance condition, followed by implementation, automated checks, a device "
        "demonstration and a short retrospective. This matches Scrum's focus on "
        "inspection and adaptation. However, Scrum roles or events should only be "
        "reported if they actually took place (Schwaber and Sutherland, 2020)."
    ),
    138: (
        "Git records the source history and GitHub hosts the MathRush3D repository. Each "
        "commit should describe one coherent change, while feature branches or pull "
        "requests are more suitable for risky work. Tags can identify the exact version "
        "submitted for assessment. Secrets such as DATABASE_URL must stay in deployment "
        "environment variables and outside the repository. GitHub Actions can run build "
        "and test workflows and retain APK files, which fits GitHub's description of "
        "Actions as a platform for customised CI/CD workflows (GitHub, 2026)."
    ),
    143: (
        "Only data needed for authentication, progress and competitive results should be "
        "collected. Leaderboard responses must never expose email addresses or password "
        "hashes. Sessions should expire, and old session records should be removable. "
        "The finished system should define how long inactive accounts are retained and "
        "provide an account-deletion process that uses the database's cascading "
        "relationships safely."
    ),
    144: (
        "The game has no chat, advertising, payments or location tracking, so several "
        "common privacy and safeguarding risks are avoided. Display names still need "
        "validation because other players can see them. Difficulty labels should be "
        "presented as game choices rather than judgements about a learner's ability, and "
        "loss messages should remain encouraging instead of embarrassing the player."
    ),
    145: (
        "Every external model and texture needs a recorded source and licence. The "
        "current repository includes local attribution information for Kenney and "
        "Quaternius assets. A replacement should only be added after its redistribution "
        "and commercial-use conditions have been checked."
    ),
    148: (
        "Testing is divided into layers so that quick, predictable checks catch rule "
        "errors before slower device and user testing begins. Unit tests cover balance, "
        "crowd formation, level progression, speed and boss-meter results. TypeScript "
        "checks cover both the browser and server code. API integration tests should use "
        "a disposable database to register a temporary account, log in, submit a run, "
        "retrieve progress and confirm leaderboard order. End-to-end tests should then "
        "cover registration, a complete level, pause and restart, next level and logout "
        "through the visible interface."
    ),
    158: (
        "At least two people should test the game independently. One should concentrate "
        "on first-time use and whether the mathematics is clear. The other should focus "
        "on mobile controls, difficulty and technical faults. Testers should rely on the "
        "instructions already present in the game rather than being coached through it. "
        "For each session, the report should record consent, device, browser or app "
        "version, date, observations, severity and suggested improvements. Assumed "
        "feedback must not be presented as if it was observed."
    ),
    163: (
        "Performance should be measured in a repeatable way rather than judged only by "
        "how smooth the game looks. Every test should record the application version, "
        "device, operating system, browser or APK version, selected level, crowd size "
        "and whether reduced effects were enabled."
    ),
    166: (
        "Every failed check should become a backlog item with reproduction steps, "
        "expected and actual results, device and build information, severity, and a "
        "supporting screenshot or recording. Feedback still needs interpretation: "
        "repeated problems, failed tasks and serious defects should carry more weight "
        "than one person's visual preference. After a correction, the original test and "
        "any related regression checks must be repeated. The final report should include "
        "at least one clear trail from a tester's observation to the decision, change, "
        "retest and evaluation."
    ),
    168: (
        "The following matrix compares the inspected project and this draft with the "
        "brief currently available. A result marked 'Partial' means that some work or "
        "evidence exists, but the criterion should not yet be claimed as fully achieved."
    ),
    180: (
        "Math Rush 3D is already a substantial application. It combines a React-based 3D "
        "game, PostgreSQL persistence, secure account handling, responsive Android "
        "delivery and automated rule tests. One of its best technical decisions is the "
        "shared generated level definition, together with common reward rules, because "
        "these reduce disagreements between scene systems and stored results. Its main "
        "weaknesses are the unresolved backend-language requirement, the missing "
        "role-based administration features and the lack of completed peer-testing and "
        "evaluation evidence."
    ),
    181: (
        "This draft covers the subjects named in the guideline supplied so far: the "
        "executive summary, SDD, problem and stakeholder analysis, requirements, success "
        "measures, risks, design diagrams, traceability, technology choices, frontend, "
        "backend and database work, methodology, version control, ethics, performance "
        "planning and peer UAT. It is a working report rather than a finished submission. "
        "The student still needs to add personal details, check every statement, include "
        "any remaining tutor guidance, carry out peer testing and complete the final "
        "evaluation."
    ),
}


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main():
    document = Document(SOURCE)
    for index, text in REWRITES.items():
        replace_paragraph_text(document.paragraphs[index], text)

    # Keep table typography aligned with the report requirement.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        rpr = run._element.get_or_add_rPr()
                        rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
                        rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                        rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Use plain academic punctuation consistently.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "\u2014" in run.text:
                run.text = run.text.replace("\u2014", ";")

    document.save(OUTPUT)
    print(f"Saved {OUTPUT}")
    print(f"Rewritten paragraphs: {len(REWRITES)}")


if __name__ == "__main__":
    main()
